import requests
import time
from docx import Document
from datetime import datetime

# requests.adapters.DEFAULT_RETRIES = 5 # 增加重连次数
# se = requests.session()
# se.keep_alive = False # 关闭多余连接

register_url = "https://www2.deepl.com/jsonrpc"
header = {
    "Content-Type": "application/json",
    "cookie": '',
    "user-agent": "",
    "authority":"www2.deepl.com",
    "method":"POST",
    "path":"/jsonrpc",
    "scheme":"https",
    "accept":"*/*",
    "accept-encoding":"gzip, deflate, br",
    "accept-language":"zh-CN,zh;q=0.9,en;q=0.8",
    "content-type":"application/json",
    "dnt":"1",
    "origin":"https://www.deepl.com",
    "referer":"https://www.deepl.com/",
    "sec-fetch-dest":"empty",
    "sec-fetch-mode":"cors",
    "sec-fetch-site":"same-site",
}
post_json = {"jsonrpc":"2.0","method": "LMT_handle_jobs","params":{"jobs":[{"kind":"default","raw_en_sentence":"have a test","raw_en_context_before":[],"raw_en_context_after":[],"preferred_num_beams":4}],"lang":{"user_preferred_langs":["ZH","EN"],"source_lang_user_selected":"auto","target_lang":"ZH"},"priority":1,"commonJobParams":{"formality":None},"timestamp":1604332586487},"id":32630018}
post_json_id = 0

doc1 = Document('test.docx')
print(str(datetime.now()) + ", 开始循环请求: ")
for idx, paragraph in enumerate(doc1.paragraphs):
    if len(paragraph.text) < 10:
        continue
    ts = int(round(time.time() * 1000))
    post_json["params"]["timestamp"] = ts
    post_json["id"] = post_json_id
    post_json_id += 1
    post_json["params"]["jobs"][0]["raw_en_sentence"] = "too many"
    print("开始请求id: " + str(post_json_id) + ", 开始时间: " + str(datetime.now()))
    print(post_json)
    response = requests.post(url=register_url, json=post_json, headers=header)
    print("请求结果: " + str(response.json()))
    print("完成请求id: " + str(post_json_id) + ", 结束时间: " + str(datetime.now()))

    doc1.paragraphs[idx].text = response.json()["result"]["translations"][0]["beams"][0]["postprocessed_sentence"]

doc1.save("test-ZH.docx")
print(str(datetime.now()) + "操作完成")

# 发送post请求
# print(datetime.now())
